#!/usr/bin/env python3
"""
🤖 Enhanced Agentic RAG System v4.0
=====================================

Production-ready RAG system with:
- Multi-format support (PDF, DOCX, TXT, PCAP)
- Secure API key entry at startup
- Modern LangChain implementation
- Enhanced error handling
- JSON-based persistence

Author: Enhanced Agentic RAG Team
Version: 4.0
"""

import os
import json
import logging
from typing import List, Dict, Tuple, Optional, Any
from dataclasses import dataclass, asdict
from datetime import datetime
from pathlib import Path
import warnings
from collections import Counter

# Suppress warnings
warnings.filterwarnings('ignore')

# LangChain imports - NEW STYLE for langchain 0.2.x with OpenAI 1.0+
from langchain.agents import initialize_agent, AgentType
from langchain.tools import Tool
from langchain_openai import ChatOpenAI, OpenAIEmbeddings
from langchain_core.documents import Document
from langchain.memory import ConversationBufferMemory
# Import text splitter - LangChain 0.1.0+ uses separate package
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS

# Document loaders
from langchain_community.document_loaders import (
    PyPDFLoader,
    TextLoader,
    UnstructuredWordDocumentLoader
)

# Tool imports
from langchain_community.utilities import WikipediaAPIWrapper
from langchain_community.tools import WikipediaQueryRun

# Tavily Search
try:
    from langchain_community.tools.tavily_search import TavilySearchResults
    TAVILY_AVAILABLE = True
except ImportError:
    TAVILY_AVAILABLE = False

# PCAP parsing
try:
    from scapy.all import rdpcap, IP, TCP, UDP
    SCAPY_AVAILABLE = True
except ImportError:
    SCAPY_AVAILABLE = False
    print("⚠️  Scapy not available. Install: pip install scapy")

# Gradio
import gradio as gr

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('agentic_rag.log'),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)


# ═══════════════════════════════════════════════════════════════════
# 📊 Data Classes
# ═══════════════════════════════════════════════════════════════════

@dataclass
class InteractionLog:
    """Log entry for each user interaction."""
    timestamp: str
    query: str
    response: str
    agent_steps: List[Dict]
    tools_used: List[str]
    feedback: Optional[str] = None
    feedback_timestamp: Optional[str] = None
    
    def to_dict(self) -> Dict:
        """Convert to dictionary for JSON serialization."""
        return asdict(self)
    
    @classmethod
    def from_dict(cls, data: Dict) -> 'InteractionLog':
        """Create from dictionary."""
        return cls(**data)


# ═══════════════════════════════════════════════════════════════════
# 🧠 Memory Manager (JSON-based, secure)
# ═══════════════════════════════════════════════════════════════════

class MemoryManager:
    """Manages conversation memory and interaction history."""
    
    def __init__(self, memory_path: str = "memory_store"):
        """Initialize memory management system."""
        self.memory_path = Path(memory_path)
        self.memory_path.mkdir(exist_ok=True)
        
        # Long-term interaction history
        self.interaction_history: List[InteractionLog] = []
        
        # Load existing memory
        self._load_memory()
        
        logger.info("🧠 Memory Manager initialized")
    
    def add_interaction(self, query: str, response: str, 
                       agent_steps: List[Dict], tools_used: List[str]):
        """Add a new interaction to memory."""
        interaction = InteractionLog(
            timestamp=datetime.now().isoformat(),
            query=query,
            response=response,
            agent_steps=agent_steps,
            tools_used=tools_used
        )
        self.interaction_history.append(interaction)
        
        # Save after every interaction
        self._save_memory()
        
        logger.info(f"💾 Interaction saved (total: {len(self.interaction_history)})")
    
    def add_feedback(self, interaction_index: int, feedback: str):
        """Add user feedback to a specific interaction."""
        if 0 <= interaction_index < len(self.interaction_history):
            self.interaction_history[interaction_index].feedback = feedback
            self.interaction_history[interaction_index].feedback_timestamp = datetime.now().isoformat()
            self._save_memory()
            logger.info(f"👍/👎 Feedback added to interaction {interaction_index}")
    
    def clear_memory(self):
        """Clear all interaction history."""
        self.interaction_history = []
        self._save_memory()
        logger.info("🗑️ All conversation history cleared")
    
    def _save_memory(self):
        """Save interaction history to disk using JSON."""
        history_file = self.memory_path / "interaction_history.json"
        try:
            with open(history_file, 'w', encoding='utf-8') as f:
                json.dump(
                    [interaction.to_dict() for interaction in self.interaction_history],
                    f,
                    indent=2,
                    ensure_ascii=False
                )
            logger.debug(f"💾 Memory saved to {history_file}")
        except Exception as e:
            logger.error(f"❌ Failed to save memory: {e}")
    
    def _load_memory(self):
        """Load interaction history from disk."""
        history_file = self.memory_path / "interaction_history.json"
        if history_file.exists():
            try:
                with open(history_file, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                    self.interaction_history = [
                        InteractionLog.from_dict(item) for item in data
                    ]
                logger.info(f"📂 Loaded {len(self.interaction_history)} past interactions")
            except Exception as e:
                logger.warning(f"⚠️ Could not load memory: {e}")
                self.interaction_history = []


# ═══════════════════════════════════════════════════════════════════
# 📄 Multi-Format Document Loader
# ═══════════════════════════════════════════════════════════════════

class MultiFormatDocumentLoader:
    """Handles loading of multiple document formats."""
    
    @staticmethod
    def load_document(file_path: str) -> Tuple[List[Document], Dict[str, Any]]:
        """
        Load document from various formats.
        
        Returns:
            Tuple of (documents, metadata)
        """
        file_path = Path(file_path)
        file_ext = file_path.suffix.lower()
        
        try:
            if file_ext == '.pdf':
                return MultiFormatDocumentLoader._load_pdf(str(file_path))
            elif file_ext in ['.docx', '.doc']:
                return MultiFormatDocumentLoader._load_docx(str(file_path))
            elif file_ext == '.txt':
                return MultiFormatDocumentLoader._load_txt(str(file_path))
            elif file_ext in ['.pcap', '.pcapng']:
                return MultiFormatDocumentLoader._load_pcap(str(file_path))
            else:
                raise ValueError(f"Unsupported file format: {file_ext}")
        except Exception as e:
            logger.error(f"Error loading document: {e}")
            raise
    
    @staticmethod
    def _load_pdf(file_path: str) -> Tuple[List[Document], Dict[str, Any]]:
        """Load PDF document."""
        logger.info(f"📄 Loading PDF: {file_path}")
        loader = PyPDFLoader(file_path)
        pages = loader.load()
        
        metadata = {
            'format': 'PDF',
            'pages': len(pages),
            'filename': Path(file_path).name
        }
        
        return pages, metadata
    
    @staticmethod
    def _load_docx(file_path: str) -> Tuple[List[Document], Dict[str, Any]]:
        """Load Word document."""
        logger.info(f"📝 Loading DOCX: {file_path}")
        
        try:
            loader = UnstructuredWordDocumentLoader(file_path)
            documents = loader.load()
        except Exception as e:
            # Fallback: try reading with python-docx directly
            logger.warning(f"Trying alternative DOCX loader: {e}")
            try:
                from docx import Document as DocxDocument
                doc = DocxDocument(file_path)
                text = '\n'.join([para.text for para in doc.paragraphs if para.text.strip()])
                documents = [Document(page_content=text, metadata={'source': file_path})]
            except Exception as e2:
                raise ValueError(f"Failed to load DOCX file: {e2}")
        
        metadata = {
            'format': 'DOCX',
            'sections': len(documents),
            'filename': Path(file_path).name
        }
        
        return documents, metadata
    
    @staticmethod
    def _load_txt(file_path: str) -> Tuple[List[Document], Dict[str, Any]]:
        """Load text file."""
        logger.info(f"📃 Loading TXT: {file_path}")
        
        try:
            loader = TextLoader(file_path, encoding='utf-8')
            documents = loader.load()
        except UnicodeDecodeError:
            # Try different encodings
            for encoding in ['latin-1', 'cp1252', 'iso-8859-1']:
                try:
                    loader = TextLoader(file_path, encoding=encoding)
                    documents = loader.load()
                    logger.info(f"Loaded with encoding: {encoding}")
                    break
                except:
                    continue
            else:
                raise ValueError("Could not decode text file with any encoding")
        
        metadata = {
            'format': 'TXT',
            'sections': len(documents),
            'filename': Path(file_path).name
        }
        
        return documents, metadata
    
    @staticmethod
    def _load_pcap(file_path: str) -> Tuple[List[Document], Dict[str, Any]]:
        """Load and analyze PCAP network capture file."""
        logger.info(f"📡 Loading PCAP: {file_path}")
        
        if not SCAPY_AVAILABLE:
            raise ImportError("Scapy is required for PCAP files. Install: pip install scapy")
        
        try:
            packets = rdpcap(file_path)
            
            # Analyze PCAP content
            analysis = []
            analysis.append(f"PCAP File Analysis: {Path(file_path).name}")
            analysis.append(f"Total Packets: {len(packets)}\n")
            
            # Protocol statistics
            protocols = {}
            ip_addresses = set()
            ports = set()
            
            for i, packet in enumerate(packets[:1000]):  # Limit to first 1000 for performance
                # Get protocol info
                if packet.haslayer(IP):
                    ip_layer = packet[IP]
                    ip_addresses.add(ip_layer.src)
                    ip_addresses.add(ip_layer.dst)
                    
                    proto = packet.sprintf("%IP.proto%")
                    protocols[proto] = protocols.get(proto, 0) + 1
                    
                    if packet.haslayer(TCP):
                        tcp_layer = packet[TCP]
                        ports.add(tcp_layer.sport)
                        ports.add(tcp_layer.dport)
                    elif packet.haslayer(UDP):
                        udp_layer = packet[UDP]
                        ports.add(udp_layer.sport)
                        ports.add(udp_layer.dport)
            
            # Build analysis text
            analysis.append("Protocol Distribution:")
            for proto, count in sorted(protocols.items(), key=lambda x: x[1], reverse=True):
                analysis.append(f"  - {proto}: {count} packets")
            
            analysis.append(f"\nUnique IP Addresses: {len(ip_addresses)}")
            analysis.append(f"Sample IPs: {', '.join(list(ip_addresses)[:10])}")
            
            analysis.append(f"\nUnique Ports: {len(ports)}")
            analysis.append(f"Sample Ports: {', '.join(map(str, sorted(list(ports))[:20]))}")
            
            # Create document
            content = '\n'.join(analysis)
            documents = [Document(page_content=content, metadata={'source': file_path, 'type': 'pcap'})]
            
            metadata = {
                'format': 'PCAP',
                'packets': len(packets),
                'protocols': len(protocols),
                'unique_ips': len(ip_addresses),
                'filename': Path(file_path).name
            }
            
            return documents, metadata
            
        except Exception as e:
            raise ValueError(f"Failed to load PCAP file: {e}")


# ═══════════════════════════════════════════════════════════════════
# 🛠️ Custom Tools
# ═══════════════════════════════════════════════════════════════════

class DocumentSearchTool:
    """Custom tool for searching uploaded documents."""
    
    def __init__(self, vectorstore: Optional[FAISS] = None):
        self.vectorstore = vectorstore
    
    def search(self, query: str) -> str:
        """Search the uploaded document for relevant information."""
        if not self.vectorstore:
            return "No document has been uploaded yet. Please upload a document first."
        
        if not query or len(query.strip()) < 2:
            return "Query too short. Please provide a more specific search term."
        
        try:
            # Search for relevant documents
            docs = self.vectorstore.similarity_search(query, k=4)
            
            if not docs:
                return "No relevant information found in the document."
            
            # Format results
            results = []
            for i, doc in enumerate(docs, 1):
                page = doc.metadata.get('page', doc.metadata.get('source', 'N/A'))
                content = doc.page_content[:300]  # First 300 chars
                results.append(f"[Source {i} - {page}]\n{content}...")
            
            return "\n\n".join(results)
        
        except Exception as e:
            logger.error(f"Error in document search: {e}")
            return f"Error searching document: {str(e)}"
    
    def update_vectorstore(self, vectorstore: FAISS):
        """Update the vectorstore when a new document is uploaded."""
        self.vectorstore = vectorstore


class PythonCalculatorTool:
    """Custom tool for mathematical calculations."""
    
    def calculate(self, expression: str) -> str:
        """Safely evaluate mathematical expressions."""
        try:
            import math
            
            # Safe namespace with math functions
            safe_dict = {
                'abs': abs, 'round': round, 'min': min, 'max': max,
                'sum': sum, 'pow': pow,
                'sqrt': math.sqrt, 'sin': math.sin, 'cos': math.cos,
                'tan': math.tan, 'log': math.log, 'exp': math.exp,
                'pi': math.pi, 'e': math.e
            }
            
            # Clean expression
            expression = expression.strip()
            
            # Evaluate expression
            result = eval(expression, {"__builtins__": {}}, safe_dict)
            return f"Result: {result}"
        
        except Exception as e:
            return f"Error calculating '{expression}': {str(e)}"


class TextAnalysisTool:
    """Custom tool for text analysis tasks."""
    
    def analyze(self, text: str) -> str:
        """Analyze text: counts words, extracts keywords, and provides summary."""
        try:
            if not text or len(text.strip()) < 3:
                return "Text too short to analyze. Please provide more text."
            
            # Word and character count
            words = text.split()
            word_count = len(words)
            char_count = len(text)
            
            # Extract keywords (frequency-based)
            meaningful_words = [w.lower() for w in words if len(w) > 4 and w.isalnum()]
            if meaningful_words:
                common = Counter(meaningful_words).most_common(5)
                keywords = [word for word, count in common]
                keywords_str = ', '.join(keywords)
            else:
                keywords_str = "None found"
            
            # Simple summary
            if len(text) > 150:
                summary = text[:100] + "..." + text[-50:]
            else:
                summary = text
            
            result = f"""TEXT ANALYSIS RESULTS:
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
📊 Statistics:
   • Words: {word_count}
   • Characters: {char_count}
   • Average word length: {char_count/word_count if word_count > 0 else 0:.1f}

🔑 Top Keywords: {keywords_str}

📝 Summary: {summary}
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━"""
            
            return result
        
        except Exception as e:
            return f"Error analyzing text: {str(e)}"


class DataFormatterTool:
    """Custom tool for data formatting and conversion."""
    
    def format(self, data: str) -> str:
        """Format data as bullet points."""
        try:
            if not data or len(data.strip()) < 2:
                return "No data provided to format."
            
            items = []
            
            # Check if already formatted
            if '•' in data or '- ' in data or data.strip().startswith(('1.', '2.', '3.')):
                return f"Already formatted:\n{data}"
            
            # Try different separators
            if ',' in data:
                items = [item.strip() for item in data.split(',') if item.strip()]
            elif '\n' in data:
                items = [item.strip() for item in data.split('\n') if item.strip()]
            elif ';' in data:
                items = [item.strip() for item in data.split(';') if item.strip()]
            else:
                if len(data.split()) > 10:
                    return f"FORMATTED TEXT:\n━━━━━━━━━━━━━━━━━━━━━━\n{data}\n━━━━━━━━━━━━━━━━━━━━━━"
                items = data.split()
            
            # Format as bullet points
            if items:
                result = "FORMATTED AS BULLET POINTS:\n━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                for item in items:
                    result += f"• {item}\n"
                result += "━━━━━━━━━━━━━━━━━━━━━━━━━━━━"
                return result
            else:
                return f"Could not parse data: {data}"
        
        except Exception as e:
            return f"Error formatting data: {str(e)}\nOriginal data: {data}"


# ═══════════════════════════════════════════════════════════════════
# 🤖 Agentic RAG System - MODERN IMPLEMENTATION
# ═══════════════════════════════════════════════════════════════════

class AgenticRAG:
    """
    Modern Agentic RAG System using create_react_agent.
    
    Features:
    - Multi-format document support (PDF, DOCX, TXT, PCAP)
    - Modern LangChain implementation
    - Enhanced error handling
    - JSON-based persistence
    """
    
    def __init__(self, api_keys: Dict[str, str], memory_path: str = "memory_store"):
        """Initialize the Agentic RAG System."""
        logger.info("🚀 Initializing Agentic RAG System...")
        
        # Set API keys
        self._set_api_keys(api_keys)
        
        # Core components
        try:
            # Use newer OpenAI API (1.0+) compatible initialization
            self.llm = ChatOpenAI(model="gpt-3.5-turbo", temperature=0, openai_api_key=api_keys.get('openai'))
            self.embeddings = OpenAIEmbeddings(openai_api_key=api_keys.get('openai'))
        except Exception as e:
            raise ValueError(f"Failed to initialize OpenAI components. Check your API key: {e}")
        
        self.vectorstore: Optional[FAISS] = None
        self.document_metadata: Dict[str, Any] = {}
        
        # Memory manager
        self.memory_manager = MemoryManager(memory_path)
        
        # Custom tool instances
        self.doc_search_tool = DocumentSearchTool()
        self.calculator_tool = PythonCalculatorTool()
        self.text_analysis_tool = TextAnalysisTool()
        self.data_formatter_tool = DataFormatterTool()
        
        # Conversation memory for agent
        self.agent_memory = ConversationBufferMemory(
            memory_key="chat_history",
            return_messages=True
        )
        
        # Agent (initialized with tools)
        self.agent_executor = None
        
        # Initialize agent
        self._initialize_agent()
        
        # Load recent conversations
        self._load_recent_conversations_to_memory()
        
        logger.info("✅ Agentic RAG System initialized")
    
    def _set_api_keys(self, api_keys: Dict[str, str]):
        """Set API keys in environment."""
        if api_keys.get('openai'):
            os.environ['OPENAI_API_KEY'] = api_keys['openai']
        if api_keys.get('tavily'):
            os.environ['TAVILY_API_KEY'] = api_keys['tavily']
    
    def _load_recent_conversations_to_memory(self):
        """Load recent conversations from MemoryManager into agent memory."""
        recent = self.memory_manager.interaction_history[-10:]
        
        for interaction in recent:
            self.agent_memory.save_context(
                {"input": interaction.query},
                {"output": interaction.response}
            )
        
        if recent:
            logger.info(f"📚 Loaded {len(recent)} past conversations into agent memory")
    
    def _initialize_agent(self):
        """Initialize agent using initialize_agent method (compatible with LangChain 0.1.0+)."""
        logger.info("🔧 Initializing agent with initialize_agent()...")
        
        # Create tool list
        tools = self._create_tools()
        
        try:
            # Use initialize_agent method (compatible with LangChain 0.1.0+)
            self.agent_executor = initialize_agent(
                tools=tools,
                llm=self.llm,
                agent=AgentType.ZERO_SHOT_REACT_DESCRIPTION,
                verbose=True,
                handle_parsing_errors=True,
                max_iterations=6,
                return_intermediate_steps=True
            )
            
            logger.info(f"✅ Agent initialized with {len(tools)} tools (ZERO_SHOT_REACT agent)")
        except Exception as e:
            logger.error(f"❌ Failed to initialize agent: {e}")
            raise
    
    def _create_tools(self) -> List[Tool]:
        """Create all available tools for the agent."""
        tools = []
        
        # 1. Document Search Tool
        tools.append(
            Tool(
                name="DocumentSearch",
                func=self.doc_search_tool.search,
                description="""Search uploaded documents (PDF, DOCX, TXT, PCAP). Use when user asks about 'the document', 'the file', or 'uploaded content'.
Input: search keywords (e.g., 'databases', 'Azure Storage').
Example: 'What does the document say about X?' → input: 'X'."""
            )
        )
        
        # 2. Calculator Tool
        tools.append(
            Tool(
                name="Calculator",
                func=self.calculator_tool.calculate,
                description="""Perform mathematical calculations.
Input: Math expression (e.g., '25*4', '100/12', 'sqrt(16)').
Example: 'Calculate 25 times 4' → input: '25*4'."""
            )
        )
        
        # 3. Text Analysis Tool
        tools.append(
            Tool(
                name="TextAnalysis",
                func=self.text_analysis_tool.analyze,
                description="""Analyze text to get word count, keywords, and summary.
Input: The text to analyze.
Example: 'Analyze this: [text]' → input: '[text]'."""
            )
        )
        
        # 4. Data Formatter Tool
        tools.append(
            Tool(
                name="DataFormatter",
                func=self.data_formatter_tool.format,
                description="""Format items as bullet points.
Input: Comma-separated items.
Example: 'Format: A, B, C' → input: 'A, B, C'."""
            )
        )
        
        # 5. Web Search Tool (Tavily) - wrapped in Tool for compatibility
        if TAVILY_AVAILABLE and os.getenv('TAVILY_API_KEY'):
            try:
                tavily_search = TavilySearchResults(max_results=3)
                tavily_tool = Tool(
                    name="WebSearch",
                    func=tavily_search.run,
                    description="""Search the internet for current information. Use ONLY when user asks for 'latest', 'today', 'now', '2024', '2025', 'current', or 'recent' info.
Input: Search query.
Example: 'Latest Azure pricing' → input: 'Azure pricing 2024'."""
                )
                tools.append(tavily_tool)
                logger.info("✅ Tavily WebSearch tool added")
            except Exception as e:
                logger.warning(f"⚠️ Tavily tool failed: {e}")
        
        # 6. Wikipedia Tool - wrapped in Tool for compatibility
        try:
            wiki_api = WikipediaQueryRun(api_wrapper=WikipediaAPIWrapper())
            wiki_tool = Tool(
                name="Wikipedia",
                func=wiki_api.run,
                description="""Search Wikipedia for factual information.
Input: Search topic.
Example: 'Who is Albert Einstein?' → input: 'Albert Einstein'."""
            )
            tools.append(wiki_tool)
        except Exception as e:
            logger.warning(f"⚠️ Wikipedia tool failed: {e}")
        
        logger.info(f"📦 Created {len(tools)} tools")
        return tools
    
    def process_document(self, file_path: str) -> Dict[str, Any]:
        """Process document of any supported format."""
        try:
            logger.info(f"📄 Processing document: {file_path}")
            
            # Load document
            documents, doc_metadata = MultiFormatDocumentLoader.load_document(file_path)
            
            if not documents:
                return {
                    'success': False,
                    'error': 'No content extracted from document'
                }
            
            # Split into chunks
            text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=1000,
                chunk_overlap=200,
                separators=["\n\n", "\n", " ", ""]
            )
            chunks = text_splitter.split_documents(documents)
            
            # Create vector store with proper embeddings
            texts = [doc.page_content for doc in chunks]
            metadatas = [doc.metadata for doc in chunks]
            self.vectorstore = FAISS.from_texts(texts, self.embeddings, metadatas=metadatas)
            
            # Update document search tool
            self.doc_search_tool.update_vectorstore(self.vectorstore)
            
            # Save to disk
            vectorstore_path = Path("faiss_index")
            vectorstore_path.mkdir(exist_ok=True)
            self.vectorstore.save_local(str(vectorstore_path))
            
            # Store metadata
            self.document_metadata = {
                **doc_metadata,
                'chunks': len(chunks),
                'timestamp': datetime.now().isoformat()
            }
            
            logger.info(f"✅ Document processed: {doc_metadata.get('format')} - {len(chunks)} chunks")
            
            return {
                'success': True,
                **self.document_metadata
            }
            
        except Exception as e:
            logger.error(f"❌ Document processing failed: {e}")
            import traceback
            traceback.print_exc()
            return {
                'success': False,
                'error': str(e)
            }
    
    def chat(self, query: str) -> Dict[str, Any]:
        """Main chat interface."""
        if not self.agent_executor:
            return {
                "response": "❌ Agent not initialized properly.",
                "metadata": {"error": "Agent initialization failed"}
            }
        
        if not query or len(query.strip()) < 2:
            return {
                "response": "Please provide a valid question.",
                "metadata": {"error": "Empty query"}
            }
        
        try:
            logger.info(f"💬 Processing query: {query[:100]}...")
            
            # Invoke agent
            result = self.agent_executor.invoke({"input": query})
            
            # Extract response
            response = result.get('output', 'No response generated')
            
            # Extract intermediate steps
            intermediate_steps = result.get('intermediate_steps', [])
            tools_used = []
            agent_steps = []
            
            for step in intermediate_steps:
                if len(step) >= 2:
                    action, observation = step[0], step[1]
                    tool_name = action.tool if hasattr(action, 'tool') else 'Unknown'
                    tool_input = action.tool_input if hasattr(action, 'tool_input') else ''
                    
                    if tool_name not in tools_used:
                        tools_used.append(tool_name)
                    
                    agent_steps.append({
                        'tool': tool_name,
                        'input': str(tool_input)[:100],
                        'output': str(observation)[:200] + '...' if len(str(observation)) > 200 else str(observation)
                    })
            
            # Save to memory
            self.memory_manager.add_interaction(
                query=query,
                response=response,
                agent_steps=agent_steps,
                tools_used=tools_used
            )
            
            # Also save to agent memory
            self.agent_memory.save_context(
                {"input": query},
                {"output": response}
            )
            
            result_dict = {
                "response": response,
                "metadata": {
                    "tools_used": tools_used,
                    "num_steps": len(agent_steps),
                    "agent_reasoning": agent_steps
                },
                "conversation_id": len(self.memory_manager.interaction_history) - 1
            }
            
            logger.info(f"✅ Query completed. Tools used: {tools_used}")
            return result_dict
            
        except Exception as e:
            logger.error(f"❌ Error in chat: {e}")
            import traceback
            traceback.print_exc()
            return {
                "response": f"❌ An error occurred: {str(e)}",
                "metadata": {"error": str(e)}
            }
    
    def add_feedback(self, conversation_id: int, feedback: str):
        """Add user feedback."""
        self.memory_manager.add_feedback(conversation_id, feedback)
    
    def get_conversation_history(self, num_interactions: int = 10) -> List[InteractionLog]:
        """Get recent conversation history."""
        return self.memory_manager.interaction_history[-num_interactions:]
    
    def clear_memory(self):
        """Clear conversation memory."""
        self.memory_manager.clear_memory()
        self.agent_memory.clear()
        logger.info("🗑️ Memory cleared")
    
    def export_logs(self, filepath: str = "interaction_logs.json") -> bool:
        """Export interaction logs to JSON."""
        try:
            logs = [log.to_dict() for log in self.memory_manager.interaction_history]
            with open(filepath, 'w', encoding='utf-8') as f:
                json.dump(logs, f, indent=2, ensure_ascii=False)
            logger.info(f"📤 Logs exported to {filepath}")
            return True
        except Exception as e:
            logger.error(f"❌ Failed to export logs: {e}")
            return False


print("✅ Core system defined successfully!")

